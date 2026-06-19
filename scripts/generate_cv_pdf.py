#!/usr/bin/env python3
"""
Generate a modern, professional PDF and DOCX CV from the YAML data file.

Requirements:
    pip install reportlab pyyaml bibtexparser python-docx pylatexenc

Usage:
    python scripts/generate_cv_pdf.py
"""

import os
import re
from datetime import datetime

import yaml
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

try:
    import bibtexparser
    BIBTEX_AVAILABLE = True
except ImportError:
    BIBTEX_AVAILABLE = False
    print("Warning: bibtexparser not installed. Publications will be skipped.")
    print("Install with: pip install bibtexparser")

try:
    from pylatexenc.latex2text import LatexNodes2Text
    LATEX_AVAILABLE = True
except ImportError:
    LATEX_AVAILABLE = False
    print("Warning: pylatexenc not installed. LaTeX accents may not render correctly.")
    print("Install with: pip install pylatexenc")

try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. DOCX generation will be skipped.")
    print("Install with: pip install python-docx")


# ----------------------------------------------------------------------------
# Shared data helpers
# ----------------------------------------------------------------------------

ACTIVITY_FILES = {
    'chair': 'activities_chair.yml',
    'pc': 'activities_pc.yml',
    'conference': 'activities_conference.yml',
    'journal': 'activities_journal.yml',
    'editor': 'activities_editor.yml',
    'funding': 'activities_funding.yml',
    'sc': 'activities_sc.yml',
}


def load_activities(project_root):
    activities_dir = os.path.join(project_root, '_data')
    out = {}
    for key, filename in ACTIVITY_FILES.items():
        filepath = os.path.join(activities_dir, filename)
        if os.path.exists(filepath):
            with open(filepath, 'r') as f:
                out[key] = yaml.safe_load(f)
    return out


def load_publications(project_root):
    if not BIBTEX_AVAILABLE:
        return []
    bib_path = os.path.join(project_root, '_bibliography', 'references.bib')
    if not os.path.exists(bib_path):
        print(f"Warning: BibTeX file not found at {bib_path}")
        return []
    try:
        with open(bib_path, 'r', encoding='utf-8') as f:
            bib_database = bibtexparser.load(f)
        pubs = sorted(
            bib_database.entries,
            key=lambda x: int(x.get('year', '0')),
            reverse=True,
        )
        return pubs
    except Exception as e:
        print(f"Error loading BibTeX file: {e}")
        return []


def strip_html(text):
    if not text:
        return ""
    text = re.sub(r'<a[^>]*>', '', text)
    text = re.sub(r'</a>', '', text)
    text = re.sub(r'<[^>]+>', '', text)
    return text.strip()


def normalize_bibtex_text(text):
    if not text:
        return ""
    cleaned = text.replace('{', '').replace('}', '')
    if LATEX_AVAILABLE:
        return LatexNodes2Text().latex_to_text(cleaned).strip()
    return cleaned.strip()


def parse_amount(amount_str):
    if not amount_str:
        return 0
    cleaned = re.sub(r'[^0-9.]', '', str(amount_str))
    try:
        return float(cleaned)
    except ValueError:
        return 0


def emphasize_numbers(text):
    if not text:
        return ""
    return re.sub(r'(\$?\d[\d,]*\+?M?)', r'<b>\1</b>', text)


# ----------------------------------------------------------------------------
# PDF Generator
# ----------------------------------------------------------------------------


class CVGenerator:
    """Modern, professional PDF CV generator."""

    # --- Page geometry -------------------------------------------------------
    PAGE_WIDTH, PAGE_HEIGHT = letter
    MARGIN_X = 0.7 * inch
    MARGIN_TOP = 0.65 * inch
    MARGIN_BOTTOM = 0.85 * inch  # extra room for footer
    CONTENT_WIDTH = PAGE_WIDTH - (2 * MARGIN_X)

    # Two-column row geometry — single source of truth so dates line up
    RIGHT_COL_WIDTH = 1.25 * inch
    LEFT_COL_WIDTH = CONTENT_WIDTH - RIGHT_COL_WIDTH

    # Header right column is wider to accommodate contact info
    HEADER_RIGHT_WIDTH = 2.6 * inch
    HEADER_LEFT_WIDTH = CONTENT_WIDTH - HEADER_RIGHT_WIDTH

    # Vertical rhythm
    SPACE_SECTION = 0.14 * inch
    SPACE_SUBSECTION = 0.07 * inch
    SPACE_ITEM = 0.04 * inch
    SPACE_DETAIL = 0.02 * inch

    # Sub-entry indentation (kept consistent across appointments, funding, etc.)
    INDENT_SUB = 0.18 * inch
    INDENT_DETAIL = 0.18 * inch

    def __init__(self, yaml_path, project_root):
        with open(yaml_path, 'r') as f:
            self.data = yaml.safe_load(f)

        self.project_root = project_root

        # Refined modern palette: deep navy primary, royal blue accent
        self.palette = {
            'ink':     colors.HexColor('#1a202c'),
            'primary': colors.HexColor('#1e3a8a'),
            'accent':  colors.HexColor('#2563eb'),
            'muted':   colors.HexColor('#64748b'),
            'border':  colors.HexColor('#e2e8f0'),
            'soft':    colors.HexColor('#f8fafc'),
            'card':    colors.HexColor('#eff6ff'),
        }
        self.styles = self._create_styles()
        self.story = []

        self.activities = load_activities(project_root)
        self.publications = load_publications(project_root)
        if self.publications:
            print(f"Loaded {len(self.publications)} publications from BibTeX file")

    # ------------------------------------------------------------------ styles
    def _create_styles(self):
        styles = getSampleStyleSheet()

        styles.add(ParagraphStyle(
            name='CVName', parent=styles['Normal'],
            fontName='Helvetica-Bold', fontSize=22, leading=26,
            textColor=self.palette['ink'], spaceAfter=2, alignment=TA_LEFT,
        ))
        styles.add(ParagraphStyle(
            name='CVTagline', parent=styles['Normal'],
            fontName='Helvetica-Bold', fontSize=11, leading=14,
            textColor=self.palette['accent'], spaceAfter=2, alignment=TA_LEFT,
        ))
        styles.add(ParagraphStyle(
            name='CVTitleMuted', parent=styles['Normal'],
            fontName='Helvetica', fontSize=9, leading=11,
            textColor=self.palette['muted'], spaceAfter=0, alignment=TA_LEFT,
        ))
        styles.add(ParagraphStyle(
            name='ContactPrimary', parent=styles['Normal'],
            fontName='Helvetica-Bold', fontSize=9, leading=11,
            textColor=self.palette['ink'], spaceAfter=1, alignment=TA_RIGHT,
        ))
        styles.add(ParagraphStyle(
            name='ContactInfo', parent=styles['Normal'],
            fontName='Helvetica', fontSize=8.5, leading=11,
            textColor=self.palette['muted'], spaceAfter=1, alignment=TA_RIGHT,
        ))
        styles.add(ParagraphStyle(
            name='ContactLink', parent=styles['Normal'],
            fontName='Helvetica', fontSize=8.5, leading=11,
            textColor=self.palette['accent'], spaceAfter=1, alignment=TA_RIGHT,
        ))

        styles.add(ParagraphStyle(
            name='SectionHeading', parent=styles['Normal'],
            fontName='Helvetica-Bold', fontSize=10.5, leading=12,
            textColor=self.palette['primary'], spaceBefore=0, spaceAfter=0,
            alignment=TA_LEFT,
        ))
        styles.add(ParagraphStyle(
            name='SubsectionHeading', parent=styles['Normal'],
            fontName='Helvetica-Bold', fontSize=10, leading=12,
            textColor=self.palette['ink'], spaceBefore=0, spaceAfter=0,
            alignment=TA_LEFT, keepWithNext=1,
        ))
        styles.add(ParagraphStyle(
            name='InstitutionHeading', parent=styles['Normal'],
            fontName='Helvetica-Bold', fontSize=10, leading=12,
            textColor=self.palette['primary'], spaceBefore=0, spaceAfter=0,
            alignment=TA_LEFT, keepWithNext=1,
        ))
        styles.add(ParagraphStyle(
            name='CategoryHeading', parent=styles['Normal'],
            fontName='Helvetica-Bold', fontSize=9, leading=11,
            textColor=self.palette['accent'], spaceBefore=0, spaceAfter=0,
            alignment=TA_LEFT, keepWithNext=1,
        ))

        styles.add(ParagraphStyle(
            name='CVBody', parent=styles['Normal'],
            fontName='Helvetica', fontSize=9.5, leading=12.5,
            textColor=self.palette['ink'], spaceAfter=2, alignment=TA_JUSTIFY,
        ))
        styles.add(ParagraphStyle(
            name='CVEntry', parent=styles['Normal'],
            fontName='Helvetica', fontSize=9.5, leading=12.5,
            textColor=self.palette['ink'], spaceAfter=0, alignment=TA_LEFT,
        ))
        styles.add(ParagraphStyle(
            name='CVSmall', parent=styles['Normal'],
            fontName='Helvetica', fontSize=8.5, leading=11,
            textColor=self.palette['muted'], spaceAfter=0, alignment=TA_LEFT,
        ))
        styles.add(ParagraphStyle(
            name='CVDetail', parent=styles['Normal'],
            fontName='Helvetica', fontSize=8.5, leading=11,
            textColor=self.palette['muted'], spaceAfter=0, alignment=TA_LEFT,
            leftIndent=self.INDENT_DETAIL,
        ))

        # Right-aligned metadata (dates, amounts)
        styles.add(ParagraphStyle(
            name='EntryMeta', parent=styles['Normal'],
            fontName='Helvetica-Bold', fontSize=8.5, leading=11,
            textColor=self.palette['accent'], spaceAfter=0, alignment=TA_RIGHT,
        ))
        styles.add(ParagraphStyle(
            name='EntryMetaMuted', parent=styles['Normal'],
            fontName='Helvetica', fontSize=8.5, leading=11,
            textColor=self.palette['muted'], spaceAfter=0, alignment=TA_RIGHT,
        ))

        # Sub-entry (inside indented column)
        styles.add(ParagraphStyle(
            name='SubEntry', parent=styles['Normal'],
            fontName='Helvetica', fontSize=9.5, leading=12.5,
            textColor=self.palette['ink'], spaceAfter=0, alignment=TA_LEFT,
        ))

        # Publication style — hanging indent so the number stays at left edge
        styles.add(ParagraphStyle(
            name='Publication', parent=styles['Normal'],
            fontName='Helvetica', fontSize=8.5, leading=11,
            textColor=self.palette['ink'], spaceAfter=3, alignment=TA_JUSTIFY,
            leftIndent=0.30 * inch, firstLineIndent=-0.30 * inch,
        ))

        # Metric (big number) and label
        styles.add(ParagraphStyle(
            name='MetricValue', parent=styles['Normal'],
            fontName='Helvetica-Bold', fontSize=18, leading=20,
            textColor=self.palette['primary'], spaceAfter=0, alignment=TA_CENTER,
        ))
        styles.add(ParagraphStyle(
            name='MetricLabel', parent=styles['Normal'],
            fontName='Helvetica', fontSize=7.5, leading=9,
            textColor=self.palette['muted'], spaceAfter=0, alignment=TA_CENTER,
        ))

        # Highlight card body
        styles.add(ParagraphStyle(
            name='HighlightBody', parent=styles['Normal'],
            fontName='Helvetica', fontSize=9, leading=12,
            textColor=self.palette['ink'], spaceAfter=0, alignment=TA_JUSTIFY,
        ))

        return styles

    # ------------------------------------------------------------------ layout
    def _section(self, title):
        """Modern section header: thick left accent bar + uppercase title + hairline rule."""
        title_para = Paragraph(title.upper(), self.styles['SectionHeading'])
        tbl = Table([[title_para]], colWidths=[self.CONTENT_WIDTH])
        tbl.setStyle(TableStyle([
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('LEFTPADDING', (0, 0), (-1, -1), 8),
            ('RIGHTPADDING', (0, 0), (-1, -1), 0),
            ('TOPPADDING', (0, 0), (-1, -1), 3),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ('LINEBEFORE', (0, 0), (0, -1), 3, self.palette['primary']),
            ('LINEBELOW', (0, 0), (-1, -1), 0.5, self.palette['border']),
        ]))
        self.story.append(Spacer(1, self.SPACE_SECTION))
        self.story.append(tbl)
        self.story.append(Spacer(1, self.SPACE_ITEM + 0.02 * inch))

    def _two_col_row(self, left, right=None, left_indent=0):
        """
        Two-column row aligned across the document.

        The table is always CONTENT_WIDTH wide and the right column is always
        RIGHT_COL_WIDTH wide. Indentation is applied via cell left-padding so
        the right edge (and date column) lines up with every other row.
        """
        left_cell = left if isinstance(left, list) else [left]
        if right is None:
            right_cell = [Spacer(1, 0)]
        else:
            right_cell = right if isinstance(right, list) else [right]

        tbl = Table(
            [[left_cell, right_cell]],
            colWidths=[self.LEFT_COL_WIDTH, self.RIGHT_COL_WIDTH],
        )
        tbl.setStyle(TableStyle([
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (0, -1), left_indent),
            ('LEFTPADDING', (1, 0), (1, -1), 4),
            ('RIGHTPADDING', (0, 0), (-1, -1), 0),
            ('TOPPADDING', (0, 0), (-1, -1), 0),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
            ('ALIGN', (1, 0), (1, -1), 'RIGHT'),
        ]))
        return tbl

    def _full_width_row(self, content, left_indent=0):
        """Full-width single-cell row with optional left indent (no right column)."""
        cell = content if isinstance(content, list) else [content]
        tbl = Table([[cell]], colWidths=[self.CONTENT_WIDTH])
        tbl.setStyle(TableStyle([
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), left_indent),
            ('RIGHTPADDING', (0, 0), (-1, -1), 0),
            ('TOPPADDING', (0, 0), (-1, -1), 0),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
        ]))
        return tbl

    def _highlight_card(self, paragraphs):
        """Soft-tinted card with an accent left bar — for executive summary highlights."""
        cell = paragraphs if isinstance(paragraphs, list) else [paragraphs]
        tbl = Table([[cell]], colWidths=[self.CONTENT_WIDTH])
        tbl.setStyle(TableStyle([
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('BACKGROUND', (0, 0), (-1, -1), self.palette['card']),
            ('LEFTPADDING', (0, 0), (-1, -1), 12),
            ('RIGHTPADDING', (0, 0), (-1, -1), 12),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('LINEBEFORE', (0, 0), (0, -1), 3, self.palette['accent']),
        ]))
        return tbl

    # ------------------------------------------------------------------ header
    def _add_header(self):
        personal = self.data.get('personal', {})

        name = personal.get('name', '')
        suffix = personal.get('title', '')
        full_name = f"{name}, {suffix}" if suffix else name

        titles = []
        for pos in personal.get('titles_full', []):
            t = pos.get('title', '')
            if pos.get('interim'):
                t += ' (Interim)'
            titles.append(t)

        left_cell = [Paragraph(full_name, self.styles['CVName'])]
        if titles:
            left_cell.append(
                Paragraph(' • '.join(titles), self.styles['CVTagline'])
            )

        office = personal.get('office', {})
        contact = personal.get('contact', {})

        right_cell = []
        if office.get('institution'):
            right_cell.append(
                Paragraph(office['institution'], self.styles['ContactPrimary'])
            )
        if contact.get('email'):
            right_cell.append(
                Paragraph(contact['email'], self.styles['ContactInfo'])
            )
        if contact.get('website'):
            site = contact['website'].replace('https://', '').replace('http://', '')
            right_cell.append(Paragraph(site, self.styles['ContactLink']))
        if contact.get('phone'):
            right_cell.append(
                Paragraph(contact['phone'], self.styles['ContactInfo'])
            )
        if not right_cell:
            right_cell = [Spacer(1, 0)]

        header = Table(
            [[left_cell, right_cell]],
            colWidths=[self.HEADER_LEFT_WIDTH, self.HEADER_RIGHT_WIDTH],
        )
        header.setStyle(TableStyle([
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), 0),
            ('RIGHTPADDING', (0, 0), (-1, -1), 0),
            ('TOPPADDING', (0, 0), (-1, -1), 0),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('ALIGN', (1, 0), (1, -1), 'RIGHT'),
            ('LINEBELOW', (0, 0), (-1, -1), 1.8, self.palette['primary']),
        ]))
        self.story.append(header)
        self.story.append(Spacer(1, self.SPACE_SUBSECTION))

    # ----------------------------------------------------------- exec summary
    def _add_executive_summary(self):
        personal = self.data.get('personal', {})
        intro = personal.get('intro', '')

        if intro:
            self.story.append(
                Paragraph(strip_html(intro), self.styles['CVBody'])
            )
            self.story.append(Spacer(1, self.SPACE_ITEM + 0.02 * inch))

        self._add_metrics_row()

        research = self.data.get('research', {})
        highlights = research.get('leadership_highlights', [])
        accomplishments = list(highlights[:4])

        memberships = personal.get('memberships', [])
        senior = [m for m in memberships if 'Senior' in m.get('level', '')]
        if senior:
            orgs = ', '.join(sorted({
                m.get('organization', '') for m in senior if m.get('organization')
            }))
            if orgs:
                accomplishments.append(f"Senior Member of {orgs}")

        if accomplishments:
            text = '; '.join(emphasize_numbers(item) for item in accomplishments) + '.'
            self.story.append(self._highlight_card(
                Paragraph(text, self.styles['HighlightBody'])
            ))
            self.story.append(Spacer(1, self.SPACE_ITEM))

        projects = self.data.get('research_projects', [])
        initiative_parts = []
        for project in projects:
            name = project.get('name', '')
            count = str(project.get('count', '')).strip()
            if name and count:
                initiative_parts.append(f"<b>{name}</b> ({count})")
            elif name:
                initiative_parts.append(f"<b>{name}</b>")
        if initiative_parts:
            line = "<b>Signature initiatives:</b> " + ' • '.join(initiative_parts) + "."
            self.story.append(Paragraph(line, self.styles['CVBody']))

    def _add_metrics_row(self):
        """KPI strip: funded $, publications, grants, awards."""
        funding = self.data.get('funding', {})
        total = funding.get('total_funding', {})
        total_amount = total.get('amount') if total else None

        funding_count = sum(
            len(funding.get(cat, [])) for cat in ('doe', 'nsf', 'darpa', 'international')
        )
        if not total_amount:
            # Fallback: sum of individual award amounts
            total = sum(
                parse_amount(a.get('amount'))
                for cat in ('doe', 'nsf', 'darpa', 'international')
                for a in funding.get(cat, [])
            )
            if total >= 1_000_000:
                total_amount = f"${total/1_000_000:.0f}M+"

        awards = self.data.get('awards', [])
        pub_count = len(self.publications) if self.publications else 0

        metrics = []
        if total_amount:
            metrics.append((total_amount, 'Funded awards'))
        if pub_count:
            metrics.append((str(pub_count), 'Publications'))
        if funding_count:
            metrics.append((str(funding_count), 'Competitive grants'))
        if awards:
            metrics.append((str(len(awards)), 'Best paper awards'))

        if not metrics:
            return

        row_cells = []
        for value, label in metrics:
            cell = [
                Paragraph(value, self.styles['MetricValue']),
                Spacer(1, 1),
                Paragraph(label.upper(), self.styles['MetricLabel']),
            ]
            row_cells.append(cell)

        col_width = self.CONTENT_WIDTH / len(metrics)
        tbl = Table([row_cells], colWidths=[col_width] * len(metrics))
        style = [
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('LEFTPADDING', (0, 0), (-1, -1), 4),
            ('RIGHTPADDING', (0, 0), (-1, -1), 4),
            ('BACKGROUND', (0, 0), (-1, -1), self.palette['soft']),
            ('BOX', (0, 0), (-1, -1), 0.4, self.palette['border']),
        ]
        # Hairline dividers between metric columns
        for i in range(1, len(metrics)):
            style.append(('LINEBEFORE', (i, 0), (i, -1), 0.4, self.palette['border']))
        tbl.setStyle(TableStyle(style))
        self.story.append(tbl)
        self.story.append(Spacer(1, self.SPACE_ITEM + 0.02 * inch))

    # --------------------------------------------------------- major funding
    def _add_major_funding(self, top_n=4):
        funding = self.data.get('funding', {})
        if not funding:
            return

        categories = [
            ('doe', 'DOE'),
            ('nsf', 'NSF'),
            ('darpa', 'DARPA'),
            ('international', 'International'),
        ]
        entries = []
        for key, label in categories:
            for a in funding.get(key, []):
                entries.append({
                    'title': a.get('title', ''),
                    'role': a.get('role', ''),
                    'period': a.get('period', ''),
                    'amount': a.get('amount', ''),
                    'amount_value': parse_amount(a.get('amount')),
                    'program': a.get('program', ''),
                    'label': label,
                })
        if not entries:
            return

        entries.sort(key=lambda x: x['amount_value'], reverse=True)
        featured = entries[:top_n]

        self._section('Major Funded Programs')

        for entry in featured:
            left = [Paragraph(f"<b>{entry['title']}</b>", self.styles['CVEntry'])]
            meta_parts = [p for p in [entry['role'], entry['program'], entry['label']] if p]
            if meta_parts:
                left.append(Paragraph(' • '.join(meta_parts), self.styles['CVSmall']))

            right = []
            if entry['amount']:
                right.append(Paragraph(entry['amount'], self.styles['EntryMeta']))
            if entry['period']:
                right.append(Paragraph(entry['period'], self.styles['EntryMetaMuted']))

            self.story.append(self._two_col_row(left, right))
            self.story.append(Spacer(1, self.SPACE_ITEM + 0.01 * inch))

    # ------------------------------------------------------ professional roles
    def _add_appointments(self):
        appointments = self.data.get('appointments', {})
        if not appointments:
            return

        self._section('Professional Appointments')

        def render_block(appt, include_dept=False):
            institution = appt.get('institution', '')
            country = appt.get('country', '')
            dept = appt.get('department', '')

            inst_line = institution
            if include_dept and dept:
                inst_line += f", {dept}"
            if country:
                inst_line += f", {country}"

            self.story.append(self._full_width_row(
                Paragraph(inst_line, self.styles['InstitutionHeading'])
            ))
            self.story.append(Spacer(1, 0.02 * inch))

            for pos in appt.get('positions', []):
                title = pos.get('title', '')
                period = pos.get('period', '')
                division = pos.get('division', '')

                pos_text = f'<font color="#2563eb">▸</font>&nbsp;&nbsp;{title}'
                if division:
                    pos_text += f", <i>{division}</i>"

                self.story.append(self._two_col_row(
                    Paragraph(pos_text, self.styles['SubEntry']),
                    Paragraph(str(period), self.styles['EntryMetaMuted']),
                    left_indent=self.INDENT_SUB,
                ))
            self.story.append(Spacer(1, self.SPACE_SUBSECTION))

        for appt in appointments.get('current', []):
            render_block(appt, include_dept=False)
        for appt in appointments.get('past', []):
            render_block(appt, include_dept=True)

    def _add_education(self):
        education = self.data.get('education', [])
        if not education:
            return
        self._section('Education')

        for edu in education:
            degree = edu.get('degree', '')
            year = edu.get('year', '')
            institution = edu.get('institution', '')
            country = edu.get('country', '')
            inst_line = f"{institution}, {country}" if country else institution

            left = [
                Paragraph(f"<b>{degree}</b>", self.styles['CVEntry']),
                Paragraph(inst_line, self.styles['CVSmall']),
            ]
            if edu.get('thesis_title'):
                left.append(Paragraph(
                    f"<i>Thesis: {edu['thesis_title']}</i>",
                    self.styles['CVSmall'],
                ))

            self.story.append(self._two_col_row(
                left,
                Paragraph(str(year), self.styles['EntryMeta']),
            ))
            self.story.append(Spacer(1, self.SPACE_ITEM + 0.02 * inch))

    def _add_research(self):
        research = self.data.get('research', {})
        if not research:
            return
        self._section('Research Focus')

        description = research.get('description', '')
        if description:
            self.story.append(
                Paragraph(strip_html(description), self.styles['CVBody'])
            )
            self.story.append(Spacer(1, self.SPACE_ITEM))

        areas = research.get('areas', [])
        if areas:
            chips = ' &nbsp;·&nbsp; '.join(f"<b>{a}</b>" for a in areas)
            self.story.append(Paragraph(
                f"<font color=\"#64748b\">Focus areas: </font>{chips}.",
                self.styles['CVBody'],
            ))

    # --------------------------------------------------------- publications
    def _format_pub(self, pub, number):
        authors = normalize_bibtex_text(pub.get('author', ''))
        title = normalize_bibtex_text(pub.get('title', ''))
        venue = normalize_bibtex_text(pub.get('booktitle') or pub.get('journal') or '')
        year = pub.get('year', 'n.d.')

        if authors:
            authors = authors.replace(
                'Ferreira da Silva, Rafael', '<b>Ferreira da Silva, Rafael</b>'
            )
            authors = authors.replace(
                'da Silva, Rafael Ferreira', '<b>da Silva, Rafael Ferreira</b>'
            )

        parts = []
        if authors:
            parts.append(authors)
        if title:
            parts.append(f'&ldquo;{title}&rdquo;')
        if venue:
            parts.append(f"<i>{venue}</i>")
        if year:
            parts.append(str(year))
        if pub.get('doi'):
            parts.append(f"DOI: {pub['doi']}")
        citation = '. '.join(parts) + '.'
        return f"<b>[{number}]</b>&nbsp;&nbsp;{citation}"

    def _add_selected_publications(self, count=10):
        if not self.publications:
            return
        self._section(f'Selected Recent Publications')

        pub_number = len(self.publications)
        for pub in self.publications[:count]:
            self.story.append(Paragraph(
                self._format_pub(pub, pub_number),
                self.styles['Publication'],
            ))
            pub_number -= 1

    def _add_publications(self):
        if not self.publications:
            return
        self._section(f'Complete Publication Record ({len(self.publications)} total)')

        current_year = None
        pub_number = len(self.publications)
        for pub in self.publications:
            year = pub.get('year', 'n.d.')
            if year != current_year:
                current_year = year
                self.story.append(Spacer(1, self.SPACE_ITEM))
                self.story.append(Paragraph(
                    str(year), self.styles['CategoryHeading']
                ))
                self.story.append(Spacer(1, 0.02 * inch))

            self.story.append(Paragraph(
                self._format_pub(pub, pub_number),
                self.styles['Publication'],
            ))
            pub_number -= 1

    # ---------------------------------------------------------------- awards
    def _add_awards(self):
        awards = self.data.get('awards', [])
        if not awards:
            return
        self._section(f'Awards & Honors')

        for award in awards:
            title = award.get('title', '')
            year = award.get('year', '')
            event = award.get('event', '')
            org = award.get('organization', '')
            paper = award.get('paper', '')

            line = f"<b>{title}</b>"
            if event:
                line += f" — {event}"
            elif org:
                line += f" — {org}"

            left = [Paragraph(line, self.styles['CVEntry'])]
            if paper:
                left.append(Paragraph(f"<i>{paper}</i>", self.styles['CVSmall']))

            self.story.append(self._two_col_row(
                left,
                Paragraph(str(year), self.styles['EntryMeta']),
            ))
            self.story.append(Spacer(1, self.SPACE_ITEM + 0.01 * inch))

    # --------------------------------------------------------- full funding
    def _add_funding(self):
        funding = self.data.get('funding', {})
        if not funding:
            return
        self._section('Funding Awards')

        total = funding.get('total_funding', {})
        if total:
            summary = (
                f"<b><font color=\"#1e3a8a\">Total: {total.get('amount', '')}</font></b>"
                f" across {total.get('projects', '')} competitively reviewed proposals."
            )
            self.story.append(Paragraph(summary, self.styles['CVBody']))
            self.story.append(Spacer(1, self.SPACE_SUBSECTION))

        groups = [
            ('doe', 'U.S. Department of Energy (DOE)', False),
            ('nsf', 'U.S. National Science Foundation (NSF)', True),
            ('darpa', 'U.S. Defense Advanced Research Projects Agency (DARPA)', False),
            ('international', 'International / Other', False),
        ]
        for key, label, with_award_ids in groups:
            self._render_funding_group(funding.get(key, []), label, with_award_ids)

    def _render_funding_group(self, awards, label, with_award_ids):
        if not awards:
            return
        self.story.append(Paragraph(label, self.styles['CategoryHeading']))
        self.story.append(Spacer(1, 0.03 * inch))

        for award in awards:
            title = award.get('title', '')
            period = award.get('period', '')
            amount = award.get('amount', '')
            role = award.get('role', '')

            left = [Paragraph(f"<b>{title}</b>", self.styles['CVEntry'])]
            detail_parts = [p for p in [role, period] if p]
            detail = ', '.join(detail_parts)
            if with_award_ids:
                ids = ', '.join(
                    f"#{a.get('id', '')}" for a in award.get('awards', []) if a.get('id')
                )
                if ids:
                    detail = f"{detail} ({ids})" if detail else ids
            if detail:
                left.append(Paragraph(detail, self.styles['CVSmall']))

            self.story.append(self._two_col_row(
                left,
                Paragraph(str(amount), self.styles['EntryMeta']) if amount else None,
                left_indent=self.INDENT_SUB,
            ))
            self.story.append(Spacer(1, self.SPACE_ITEM))
        self.story.append(Spacer(1, self.SPACE_SUBSECTION))

    # -------------------------------------------------- professional activities
    def _add_professional_activities(self):
        activities = self.data.get('professional_activities', {})
        if not self.activities and not activities:
            return
        self._section('Professional Activities')

        steering = activities.get('steering_committees', [])
        if steering:
            self.story.append(Paragraph('Steering Committees', self.styles['CategoryHeading']))
            self.story.append(Spacer(1, 0.03 * inch))
            for item in steering:
                org = item.get('organization', '')
                role = item.get('role', '')
                period = item.get('period', '')
                self.story.append(self._two_col_row(
                    Paragraph(f"{role}, <b>{org}</b>", self.styles['CVEntry']),
                    Paragraph(str(period), self.styles['EntryMetaMuted']),
                    left_indent=self.INDENT_SUB,
                ))
                self.story.append(Spacer(1, self.SPACE_DETAIL))
            self.story.append(Spacer(1, self.SPACE_SUBSECTION))

        advisory = activities.get('advisory_boards', [])
        if advisory:
            self.story.append(Paragraph('Advisory Boards', self.styles['CategoryHeading']))
            self.story.append(Spacer(1, 0.03 * inch))
            for item in advisory:
                name = item.get('name', '')
                board_type = item.get('type', '')
                period = item.get('period', '')
                note = item.get('note', '')
                text = f"<b>{name}</b> ({board_type})" if board_type else f"<b>{name}</b>"
                if note:
                    text += f" — {note}"
                self.story.append(self._two_col_row(
                    Paragraph(text, self.styles['CVEntry']),
                    Paragraph(str(period), self.styles['EntryMetaMuted']),
                    left_indent=self.INDENT_SUB,
                ))
                self.story.append(Spacer(1, self.SPACE_DETAIL))
            self.story.append(Spacer(1, self.SPACE_SUBSECTION))

        chair_data = self.activities.get('chair', [])
        if chair_data:
            self.story.append(Paragraph('Conference / Workshop Chair Roles', self.styles['CategoryHeading']))
            self.story.append(Spacer(1, 0.03 * inch))
            for conf in chair_data:
                conf_name = conf.get('conference', '')
                series = conf.get('series', '')
                for entry in conf.get('entries', []):
                    role = entry.get('role', '')
                    year = entry.get('year', '')
                    location = entry.get('location', '')
                    text = f"{role}, <b>{conf_name}</b>"
                    if series:
                        text += f" ({series})"
                    if location:
                        text += f" — {location}"
                    self.story.append(self._two_col_row(
                        Paragraph(text, self.styles['CVEntry']),
                        Paragraph(str(year), self.styles['EntryMetaMuted']),
                        left_indent=self.INDENT_SUB,
                    ))
                    self.story.append(Spacer(1, self.SPACE_DETAIL))
            self.story.append(Spacer(1, self.SPACE_SUBSECTION))

        pc_data = self.activities.get('pc', [])
        if pc_data:
            self.story.append(Paragraph('Program Committee Member', self.styles['CategoryHeading']))
            self.story.append(Spacer(1, 0.03 * inch))
            for conf in pc_data:
                conf_name = conf.get('conference', '')
                series = conf.get('series', '')
                entries = conf.get('entries', [])
                years_str = ', '.join(str(e.get('year', '')) for e in entries)
                text = f"<b>{conf_name}</b>"
                if series:
                    text += f" ({series})"
                if years_str:
                    text += f" — {years_str}"
                self.story.append(self._full_width_row(
                    Paragraph(text, self.styles['CVEntry']),
                    left_indent=self.INDENT_SUB,
                ))
                self.story.append(Spacer(1, self.SPACE_DETAIL))
            self.story.append(Spacer(1, self.SPACE_SUBSECTION))

        editorial = activities.get('editorial', [])
        if editorial:
            self.story.append(Paragraph('Editorial Positions', self.styles['CategoryHeading']))
            self.story.append(Spacer(1, 0.03 * inch))
            for item in editorial:
                journal = item.get('journal', '')
                role = item.get('role', '')
                period = item.get('period', '')
                self.story.append(self._two_col_row(
                    Paragraph(f"{role}, <i>{journal}</i>", self.styles['CVEntry']),
                    Paragraph(str(period), self.styles['EntryMetaMuted']),
                    left_indent=self.INDENT_SUB,
                ))
                self.story.append(Spacer(1, self.SPACE_DETAIL))
            self.story.append(Spacer(1, self.SPACE_SUBSECTION))

        funding_rev = activities.get('funding_reviewer', [])
        if funding_rev:
            self.story.append(Paragraph('Funding Review Panels', self.styles['CategoryHeading']))
            self.story.append(Spacer(1, 0.03 * inch))
            for item in funding_rev:
                agency = item.get('agency', '')
                role = item.get('role', '')
                period = item.get('period', '')
                self.story.append(self._two_col_row(
                    Paragraph(f"{role}, <b>{agency}</b>", self.styles['CVEntry']),
                    Paragraph(str(period), self.styles['EntryMetaMuted']),
                    left_indent=self.INDENT_SUB,
                ))
                self.story.append(Spacer(1, self.SPACE_DETAIL))

    # ----------------------------------------------------------- invited talks
    def _add_invited_talks(self):
        talks = self.data.get('invited_talks', [])
        if not talks:
            return
        self._section('Invited Talks')

        for year_group in talks:
            year = year_group.get('year', '')
            entries = year_group.get('entries', [])
            self.story.append(Paragraph(str(year), self.styles['CategoryHeading']))
            self.story.append(Spacer(1, 0.03 * inch))
            for talk in entries:
                title = talk.get('title', '')
                event = talk.get('event', '')
                location = talk.get('location', '')
                self.story.append(self._full_width_row(
                    Paragraph(f"<b>{title}</b>", self.styles['CVEntry']),
                    left_indent=self.INDENT_SUB,
                ))
                detail = ', '.join(p for p in [event, location] if p)
                if detail:
                    self.story.append(self._full_width_row(
                        Paragraph(detail, self.styles['CVSmall']),
                        left_indent=self.INDENT_SUB,
                    ))
                self.story.append(Spacer(1, self.SPACE_DETAIL))
            self.story.append(Spacer(1, self.SPACE_SUBSECTION))

    # ----------------------------------------------------------- teaching
    def _add_teaching(self):
        teaching = self.data.get('teaching', [])
        if not teaching:
            return
        self._section('Teaching Experience')

        for course in teaching:
            institution = course.get('institution', '')
            course_name = course.get('course', '')
            level = course.get('level', '')
            semester = course.get('semester', '')

            left = [
                Paragraph(f"<b>{course_name}</b> &mdash; <font color=\"#64748b\">{level}</font>",
                          self.styles['CVEntry']),
                Paragraph(institution, self.styles['CVSmall']),
            ]
            self.story.append(self._two_col_row(
                left,
                Paragraph(str(semester), self.styles['EntryMetaMuted']),
            ))
            self.story.append(Spacer(1, self.SPACE_ITEM))

    # ----------------------------------------------------------- students
    def _add_students(self):
        students = self.data.get('students', {})
        if not students:
            return
        self._section('Students & Mentoring')

        sections = [
            ('phd', 'PhD Students'),
            ('worker', 'Student Workers / Interns'),
            ('dr', 'Directed Research Students'),
            ('thesis_committee', 'Thesis Committee Member'),
        ]
        for key, label in sections:
            items = students.get(key, [])
            if not items:
                continue
            self.story.append(Paragraph(label, self.styles['CategoryHeading']))
            self.story.append(Spacer(1, 0.03 * inch))
            for s in items:
                name = s.get('name', '')
                degree = s.get('degree', '')
                period = s.get('period', '') or s.get('year', '')
                institution = s.get('institution', '')
                left_text = f"<b>{name}</b>"
                meta = ', '.join(p for p in [degree, institution] if p)
                if meta:
                    left_text += f" — {meta}"
                self.story.append(self._two_col_row(
                    Paragraph(left_text, self.styles['CVEntry']),
                    Paragraph(str(period), self.styles['EntryMetaMuted']) if period else None,
                    left_indent=self.INDENT_SUB,
                ))
                self.story.append(Spacer(1, self.SPACE_DETAIL))
            self.story.append(Spacer(1, self.SPACE_SUBSECTION))

    # ----------------------------------------------------------- affiliations
    def _add_affiliations(self):
        personal = self.data.get('personal', {})
        memberships = personal.get('memberships', [])
        certifications = personal.get('certifications', [])
        if not memberships and not certifications:
            return
        self._section('Affiliations & Certifications')

        if memberships:
            self.story.append(Paragraph('Professional Memberships', self.styles['CategoryHeading']))
            self.story.append(Spacer(1, 0.03 * inch))
            for m in memberships:
                org = m.get('organization', '')
                level = m.get('level', '')
                period = m.get('period', '')
                title = f"<b>{org}</b>"
                if level:
                    title += f" — {level}"
                self.story.append(self._two_col_row(
                    Paragraph(title, self.styles['CVEntry']),
                    Paragraph(str(period), self.styles['EntryMetaMuted']),
                    left_indent=self.INDENT_SUB,
                ))
                self.story.append(Spacer(1, self.SPACE_DETAIL))
            self.story.append(Spacer(1, self.SPACE_SUBSECTION))

        if certifications:
            self.story.append(Paragraph('Certifications', self.styles['CategoryHeading']))
            self.story.append(Spacer(1, 0.03 * inch))
            for c in certifications:
                title = c.get('title', '')
                org = c.get('organization', '')
                date = c.get('date', '')
                left_text = f"<b>{title}</b>"
                if org:
                    left_text += f" — {org}"
                self.story.append(self._two_col_row(
                    Paragraph(left_text, self.styles['CVEntry']),
                    Paragraph(str(date), self.styles['EntryMetaMuted']),
                    left_indent=self.INDENT_SUB,
                ))
                self.story.append(Spacer(1, self.SPACE_DETAIL))

    # ----------------------------------------------------------- page footer
    def _draw_footer(self, canvas, _doc):
        canvas.saveState()
        # Hairline above footer
        canvas.setStrokeColor(self.palette['border'])
        canvas.setLineWidth(0.4)
        canvas.line(
            self.MARGIN_X, 0.55 * inch,
            self.PAGE_WIDTH - self.MARGIN_X, 0.55 * inch,
        )
        # Name left, page number right
        canvas.setFont('Helvetica', 8)
        canvas.setFillColor(self.palette['muted'])
        name = self.data.get('personal', {}).get('name', '')
        suffix = self.data.get('personal', {}).get('title', '')
        full = f"{name}, {suffix}" if suffix else name
        canvas.drawString(self.MARGIN_X, 0.38 * inch, full)
        # Center: updated date
        updated = datetime.now().strftime("Curriculum Vitae · %B %Y")
        canvas.drawCentredString(self.PAGE_WIDTH / 2, 0.38 * inch, updated)
        # Page number
        page_num = canvas.getPageNumber()
        canvas.setFont('Helvetica-Bold', 8)
        canvas.setFillColor(self.palette['primary'])
        canvas.drawRightString(
            self.PAGE_WIDTH - self.MARGIN_X, 0.38 * inch,
            f"Page {page_num}",
        )
        canvas.restoreState()

    # ------------------------------------------------------------------ build
    def generate(self, output_path):
        os.makedirs(os.path.dirname(output_path), exist_ok=True)

        doc = SimpleDocTemplate(
            output_path,
            pagesize=letter,
            leftMargin=self.MARGIN_X,
            rightMargin=self.MARGIN_X,
            topMargin=self.MARGIN_TOP,
            bottomMargin=self.MARGIN_BOTTOM,
            title="Curriculum Vitae — Rafael Ferreira da Silva",
            author=self.data.get('personal', {}).get('name', ''),
        )

        print("Building CV sections...")

        # Page 1: identity + at-a-glance metrics
        self._add_header()
        self._add_executive_summary()

        # Page 1-2: career core
        self._add_major_funding()
        self._add_appointments()
        self._add_education()
        self._add_research()

        # Page 2-3: recent impact
        self._add_selected_publications(count=10)
        self._add_awards()
        self._add_funding()

        # Service & teaching
        self._add_professional_activities()
        self._add_invited_talks()
        self._add_teaching()
        self._add_affiliations()
        self._add_students()

        # Page break before the long publication list
        self.story.append(PageBreak())
        self._add_publications()

        print("Generating PDF...")
        doc.build(
            self.story,
            onFirstPage=self._draw_footer,
            onLaterPages=self._draw_footer,
        )
        print(f"✓ PDF CV generated: {output_path}")


# ----------------------------------------------------------------------------
# DOCX Generator
# ----------------------------------------------------------------------------


class CVDocxGenerator:
    """Modern DOCX CV — uses real Word borders, not underscore separators."""

    PALETTE = {
        'ink':     RGBColor(0x1a, 0x20, 0x2c),
        'primary': RGBColor(0x1e, 0x3a, 0x8a),
        'accent':  RGBColor(0x25, 0x63, 0xeb),
        'muted':   RGBColor(0x64, 0x74, 0x8b),
    }

    def __init__(self, yaml_path, project_root):
        with open(yaml_path, 'r') as f:
            self.data = yaml.safe_load(f)
        self.project_root = project_root
        self.doc = Document()
        self._set_page_margins()
        self._setup_styles()
        self.activities = load_activities(project_root)
        self.publications = load_publications(project_root)

    def _set_page_margins(self):
        for section in self.doc.sections:
            section.top_margin = Inches(0.65)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(0.7)
            section.right_margin = Inches(0.7)

    def _setup_styles(self):
        styles = self.doc.styles
        existing = {s.name for s in styles}

        normal = styles['Normal']
        normal.font.name = 'Calibri'
        normal.font.size = Pt(10)
        normal.font.color.rgb = self.PALETTE['ink']

        def add(name, size, bold=False, color=None, italic=False, space_before=0, space_after=0):
            if name in existing:
                return
            s = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            s.font.name = 'Calibri'
            s.font.size = Pt(size)
            s.font.bold = bold
            s.font.italic = italic
            if color is not None:
                s.font.color.rgb = color
            s.paragraph_format.space_before = Pt(space_before)
            s.paragraph_format.space_after = Pt(space_after)
            return s

        add('CV Name', 22, bold=True, color=self.PALETTE['ink'])
        add('CV Tagline', 11, bold=True, color=self.PALETTE['accent'])
        add('CV Contact', 9, color=self.PALETTE['muted'])
        add('CV Section', 11, bold=True, color=self.PALETTE['primary'], space_before=10, space_after=4)
        add('CV Subsection', 10, bold=True, color=self.PALETTE['primary'])
        add('CV Category', 9.5, bold=True, color=self.PALETTE['accent'], space_before=4, space_after=2)
        add('CV Body', 10, color=self.PALETTE['ink'])
        add('CV Small', 9, color=self.PALETTE['muted'])

        pub_name = 'CV Publication'
        if pub_name not in existing:
            pub_style = styles.add_style(pub_name, WD_STYLE_TYPE.PARAGRAPH)
            pub_style.font.name = 'Calibri'
            pub_style.font.size = Pt(9)
            pub_style.font.color.rgb = self.PALETTE['ink']
            pub_style.paragraph_format.left_indent = Inches(0.25)
            pub_style.paragraph_format.first_line_indent = Inches(-0.25)
            pub_style.paragraph_format.space_after = Pt(2)

    # ----- helpers
    def _bottom_border(self, paragraph, size=8, color='1e3a8a'):
        """Add a colored bottom border to a paragraph (Word native, not underscores)."""
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), str(size))
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), color)
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _left_border(self, paragraph, size=24, color='1e3a8a'):
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), str(size))
        left.set(qn('w:space'), '8')
        left.set(qn('w:color'), color)
        pBdr.append(left)
        pPr.append(pBdr)

    def _section(self, title):
        p = self.doc.add_paragraph(style='CV Section')
        p.add_run(title.upper())
        self._bottom_border(p, size=8, color='1e3a8a')

    def _add_run(self, paragraph, text, bold=False, italic=False, color=None, size=None):
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic
        if color is not None:
            run.font.color.rgb = color
        if size is not None:
            run.font.size = Pt(size)
        return run

    # ----- sections
    def _add_header(self):
        personal = self.data.get('personal', {})
        name = personal.get('name', '')
        suffix = personal.get('title', '')
        full_name = f"{name}, {suffix}" if suffix else name

        p = self.doc.add_paragraph(style='CV Name')
        p.add_run(full_name)

        titles = []
        for pos in personal.get('titles_full', []):
            t = pos.get('title', '')
            if pos.get('interim'):
                t += ' (Interim)'
            titles.append(t)
        if titles:
            p = self.doc.add_paragraph(style='CV Tagline')
            p.add_run(' • '.join(titles))

        contact = personal.get('contact', {})
        office = personal.get('office', {})
        contact_parts = []
        if office.get('institution'):
            contact_parts.append(office['institution'])
        if contact.get('email'):
            contact_parts.append(contact['email'])
        if contact.get('website'):
            contact_parts.append(contact['website'].replace('https://', ''))
        if contact.get('phone'):
            contact_parts.append(contact['phone'])
        if contact_parts:
            p = self.doc.add_paragraph(style='CV Contact')
            p.add_run(' • '.join(contact_parts))
            self._bottom_border(p, size=12, color='1e3a8a')

    def _add_executive_summary(self):
        personal = self.data.get('personal', {})
        intro = personal.get('intro', '')
        if intro:
            self.doc.add_paragraph(strip_html(intro), style='CV Body')

        research = self.data.get('research', {})
        highlights = research.get('leadership_highlights', [])
        accomplishments = list(highlights[:4])

        funding = self.data.get('funding', {})
        funding_count = sum(len(funding.get(c, [])) for c in ('doe', 'nsf', 'darpa', 'international'))
        if funding_count:
            accomplishments.append(f"{funding_count} competitively reviewed grants led or co-led")

        awards = self.data.get('awards', [])
        if awards:
            accomplishments.append(f"{len(awards)} best paper awards at international conferences")

        if self.publications:
            recent = [p for p in self.publications if int(p.get('year', '0')) >= 2020]
            accomplishments.append(
                f"{len(self.publications)} peer-reviewed publications ({len(recent)} since 2020)"
            )

        senior = [m for m in personal.get('memberships', []) if 'Senior' in m.get('level', '')]
        if senior:
            orgs = ', '.join(sorted({m.get('organization', '') for m in senior if m.get('organization')}))
            if orgs:
                accomplishments.append(f"Senior Member of {orgs}")

        if accomplishments:
            p = self.doc.add_paragraph(style='CV Body')
            p.add_run('; '.join(accomplishments) + '.')
            self._left_border(p, size=24, color='2563eb')

        projects = self.data.get('research_projects', [])
        initiatives = []
        for project in projects:
            n = project.get('name', '')
            count = str(project.get('count', '')).strip()
            initiatives.append(f"{n} ({count})" if n and count else n)
        if initiatives:
            p = self.doc.add_paragraph(style='CV Body')
            self._add_run(p, 'Signature initiatives: ', bold=True)
            p.add_run(' • '.join(initiatives) + '.')

    def _add_major_funding(self, top_n=4):
        funding = self.data.get('funding', {})
        if not funding:
            return
        categories = [('doe', 'DOE'), ('nsf', 'NSF'), ('darpa', 'DARPA'), ('international', 'International')]
        entries = []
        for key, label in categories:
            for a in funding.get(key, []):
                entries.append({
                    'title': a.get('title', ''),
                    'role': a.get('role', ''),
                    'period': a.get('period', ''),
                    'amount': a.get('amount', ''),
                    'amount_value': parse_amount(a.get('amount')),
                    'program': a.get('program', ''),
                    'label': label,
                })
        if not entries:
            return
        entries.sort(key=lambda x: x['amount_value'], reverse=True)
        self._section('Major Funded Programs')

        for entry in entries[:top_n]:
            p = self.doc.add_paragraph(style='CV Body')
            self._add_run(p, entry['title'], bold=True)
            if entry['amount']:
                self._add_run(p, f" — {entry['amount']}", bold=True, color=self.PALETTE['accent'])
            meta = ' • '.join(x for x in [entry['role'], entry['program'], entry['label'], entry['period']] if x)
            if meta:
                self.doc.add_paragraph(meta, style='CV Small')

    def _add_appointments(self):
        appointments = self.data.get('appointments', {})
        if not appointments:
            return
        self._section('Professional Appointments')

        def render(appt, include_dept=False):
            institution = appt.get('institution', '')
            country = appt.get('country', '')
            dept = appt.get('department', '')
            line = institution
            if include_dept and dept:
                line += f", {dept}"
            if country:
                line += f", {country}"
            p = self.doc.add_paragraph(style='CV Subsection')
            p.add_run(line)
            for pos in appt.get('positions', []):
                title = pos.get('title', '')
                period = pos.get('period', '')
                division = pos.get('division', '')
                pp = self.doc.add_paragraph(style='CV Body')
                self._add_run(pp, '▸  ', color=self.PALETTE['accent'])
                pp.add_run(title)
                if division:
                    self._add_run(pp, f", {division}", italic=True)
                if period:
                    self._add_run(pp, f"  ({period})", color=self.PALETTE['muted'])

        for appt in appointments.get('current', []):
            render(appt, include_dept=False)
        for appt in appointments.get('past', []):
            render(appt, include_dept=True)

    def _add_education(self):
        education = self.data.get('education', [])
        if not education:
            return
        self._section('Education')
        for edu in education:
            degree = edu.get('degree', '')
            year = edu.get('year', '')
            institution = edu.get('institution', '')
            country = edu.get('country', '')
            p = self.doc.add_paragraph(style='CV Subsection')
            p.add_run(degree)
            if year:
                self._add_run(p, f"  ({year})", color=self.PALETTE['muted'])
            inst_line = f"{institution}, {country}" if country else institution
            self.doc.add_paragraph(inst_line, style='CV Body')
            if edu.get('thesis_title'):
                p = self.doc.add_paragraph(style='CV Small')
                run = p.add_run(f"Thesis: {edu['thesis_title']}")
                run.italic = True

    def _add_research(self):
        research = self.data.get('research', {})
        if not research:
            return
        self._section('Research Focus')
        description = research.get('description', '')
        if description:
            self.doc.add_paragraph(strip_html(description), style='CV Body')
        areas = research.get('areas', [])
        if areas:
            p = self.doc.add_paragraph(style='CV Body')
            self._add_run(p, 'Focus areas: ', color=self.PALETTE['muted'])
            self._add_run(p, '  ·  '.join(areas) + '.', bold=True)

    def _format_pub(self, pub, number):
        authors = normalize_bibtex_text(pub.get('author', ''))
        title = normalize_bibtex_text(pub.get('title', ''))
        venue = normalize_bibtex_text(pub.get('booktitle') or pub.get('journal') or '')
        year = pub.get('year', 'n.d.')
        parts = []
        if authors:
            parts.append(authors)
        if title:
            parts.append(f'"{title}"')
        if venue:
            parts.append(venue)
        if year:
            parts.append(str(year))
        if pub.get('doi'):
            parts.append(f"DOI: {pub['doi']}")
        return f"[{number}]  " + '. '.join(parts) + '.'

    def _add_selected_publications(self, count=10):
        if not self.publications:
            return
        self._section('Selected Recent Publications')
        n = len(self.publications)
        for pub in self.publications[:count]:
            self.doc.add_paragraph(self._format_pub(pub, n), style='CV Publication')
            n -= 1

    def _add_publications(self):
        if not self.publications:
            return
        self._section(f'Complete Publication Record ({len(self.publications)} total)')
        current_year = None
        n = len(self.publications)
        for pub in self.publications:
            year = pub.get('year', 'n.d.')
            if year != current_year:
                current_year = year
                p = self.doc.add_paragraph(style='CV Category')
                p.add_run(str(year))
            self.doc.add_paragraph(self._format_pub(pub, n), style='CV Publication')
            n -= 1

    def _add_funding(self):
        funding = self.data.get('funding', {})
        if not funding:
            return
        self._section('Funding Awards')
        total = funding.get('total_funding', {})
        if total:
            p = self.doc.add_paragraph(style='CV Body')
            self._add_run(p, f"Total: {total.get('amount', '')}", bold=True, color=self.PALETTE['primary'])
            p.add_run(f" across {total.get('projects', '')} competitively reviewed proposals.")

        groups = [
            ('doe', 'U.S. Department of Energy (DOE)'),
            ('nsf', 'U.S. National Science Foundation (NSF)'),
            ('darpa', 'U.S. Defense Advanced Research Projects Agency (DARPA)'),
            ('international', 'International / Other'),
        ]
        for key, label in groups:
            self._render_funding_group(funding.get(key, []), label, with_award_ids=(key == 'nsf'))

    def _render_funding_group(self, awards, label, with_award_ids=False):
        if not awards:
            return
        p = self.doc.add_paragraph(style='CV Category')
        p.add_run(label)
        for a in awards:
            p = self.doc.add_paragraph(style='CV Body')
            self._add_run(p, a.get('title', ''), bold=True)
            if a.get('amount'):
                self._add_run(p, f" — {a['amount']}", bold=True, color=self.PALETTE['accent'])
            detail_parts = [x for x in [a.get('role'), a.get('period')] if x]
            detail = ', '.join(detail_parts)
            if with_award_ids:
                ids = ', '.join(f"#{x.get('id', '')}" for x in a.get('awards', []) if x.get('id'))
                if ids:
                    detail = f"{detail} ({ids})" if detail else ids
            if detail:
                self.doc.add_paragraph(detail, style='CV Small')

    def _add_awards(self):
        awards = self.data.get('awards', [])
        if not awards:
            return
        self._section('Awards & Honors')
        for award in awards:
            p = self.doc.add_paragraph(style='CV Body')
            self._add_run(p, award.get('title', ''), bold=True)
            event = award.get('event') or award.get('organization', '')
            if event:
                p.add_run(f" — {event}")
            if award.get('year'):
                self._add_run(p, f"  ({award['year']})", color=self.PALETTE['muted'])
            if award.get('paper'):
                p = self.doc.add_paragraph(style='CV Small')
                p.add_run(award['paper']).italic = True

    def _add_professional_activities(self):
        activities = self.data.get('professional_activities', {})
        if not self.activities and not activities:
            return
        self._section('Professional Activities')

        def write_category(name, items, fmt):
            if not items:
                return
            p = self.doc.add_paragraph(style='CV Category')
            p.add_run(name)
            for item in items:
                self.doc.add_paragraph(fmt(item), style='CV Body')

        write_category(
            'Steering Committees',
            activities.get('steering_committees', []),
            lambda i: f"{i.get('role', '')}, {i.get('organization', '')} ({i.get('period', '')})",
        )
        write_category(
            'Advisory Boards',
            activities.get('advisory_boards', []),
            lambda i: f"{i.get('name', '')} ({i.get('type', '')}){' — ' + i.get('note', '') if i.get('note') else ''} ({i.get('period', '')})",
        )

        chair = self.activities.get('chair', [])
        if chair:
            p = self.doc.add_paragraph(style='CV Category')
            p.add_run('Conference / Workshop Chair Roles')
            for conf in chair:
                conf_name = conf.get('conference', '')
                series = conf.get('series', '')
                for e in conf.get('entries', []):
                    text = f"{e.get('role', '')}, {conf_name}"
                    if series:
                        text += f" ({series})"
                    text += f", {e.get('location', '')}, {e.get('year', '')}"
                    self.doc.add_paragraph(text, style='CV Body')

        pc = self.activities.get('pc', [])
        if pc:
            p = self.doc.add_paragraph(style='CV Category')
            p.add_run('Program Committee Member')
            for conf in pc:
                conf_name = conf.get('conference', '')
                series = conf.get('series', '')
                years = ', '.join(str(e.get('year', '')) for e in conf.get('entries', []))
                text = conf_name
                if series:
                    text += f" ({series})"
                text += f": {years}"
                self.doc.add_paragraph(text, style='CV Body')

        write_category(
            'Editorial Positions',
            activities.get('editorial', []),
            lambda i: f"{i.get('role', '')}, {i.get('journal', '')} ({i.get('period', '')})",
        )
        write_category(
            'Funding Review Panels',
            activities.get('funding_reviewer', []),
            lambda i: f"{i.get('role', '')}, {i.get('agency', '')} ({i.get('period', '')})",
        )

    def _add_invited_talks(self):
        talks = self.data.get('invited_talks', [])
        if not talks:
            return
        self._section('Invited Talks')
        for year_group in talks:
            p = self.doc.add_paragraph(style='CV Category')
            p.add_run(str(year_group.get('year', '')))
            for talk in year_group.get('entries', []):
                p = self.doc.add_paragraph(style='CV Body')
                p.add_run(talk.get('title', '')).bold = True
                detail = ', '.join(x for x in [talk.get('event', ''), talk.get('location', '')] if x)
                if detail:
                    self.doc.add_paragraph(detail, style='CV Small')

    def _add_teaching(self):
        teaching = self.data.get('teaching', [])
        if not teaching:
            return
        self._section('Teaching Experience')
        for course in teaching:
            p = self.doc.add_paragraph(style='CV Body')
            self._add_run(p, f"{course.get('course', '')}", bold=True)
            level = course.get('level', '')
            if level:
                self._add_run(p, f" — {level}", color=self.PALETTE['muted'])
            semester = course.get('semester', '')
            if semester:
                self._add_run(p, f"  ({semester})", color=self.PALETTE['muted'])
            inst = course.get('institution', '')
            if inst:
                self.doc.add_paragraph(inst, style='CV Small')

    def _add_students(self):
        students = self.data.get('students', {})
        if not students:
            return
        self._section('Students & Mentoring')
        sections = [
            ('phd', 'PhD Students'),
            ('worker', 'Student Workers / Interns'),
            ('dr', 'Directed Research Students'),
            ('thesis_committee', 'Thesis Committee Member'),
        ]
        for key, label in sections:
            items = students.get(key, [])
            if not items:
                continue
            p = self.doc.add_paragraph(style='CV Category')
            p.add_run(label)
            for s in items:
                name = s.get('name', '')
                degree = s.get('degree', '')
                period = s.get('period', '') or s.get('year', '')
                institution = s.get('institution', '')
                p = self.doc.add_paragraph(style='CV Body')
                self._add_run(p, name, bold=True)
                trail = ', '.join(x for x in [degree, institution] if x)
                if trail:
                    p.add_run(f" — {trail}")
                if period:
                    self._add_run(p, f" ({period})", color=self.PALETTE['muted'])

    def _add_affiliations(self):
        personal = self.data.get('personal', {})
        memberships = personal.get('memberships', [])
        certifications = personal.get('certifications', [])
        if not memberships and not certifications:
            return
        self._section('Affiliations & Certifications')

        if memberships:
            p = self.doc.add_paragraph(style='CV Category')
            p.add_run('Professional Memberships')
            for m in memberships:
                p = self.doc.add_paragraph(style='CV Body')
                self._add_run(p, m.get('organization', ''), bold=True)
                if m.get('level'):
                    p.add_run(f" — {m['level']}")
                if m.get('period'):
                    self._add_run(p, f" ({m['period']})", color=self.PALETTE['muted'])

        if certifications:
            p = self.doc.add_paragraph(style='CV Category')
            p.add_run('Certifications')
            for c in certifications:
                p = self.doc.add_paragraph(style='CV Body')
                self._add_run(p, c.get('title', ''), bold=True)
                if c.get('organization'):
                    p.add_run(f" — {c['organization']}")
                if c.get('date'):
                    self._add_run(p, f" ({c['date']})", color=self.PALETTE['muted'])

    def _add_footer(self):
        p = self.doc.add_paragraph(style='CV Small')
        self._bottom_border(p, size=4, color='e2e8f0')
        updated = datetime.now().strftime("Last updated: %B %Y")
        p = self.doc.add_paragraph(style='CV Small')
        p.add_run(updated)

    def generate(self, output_path):
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        print("Building DOCX CV sections...")
        self._add_header()
        self._add_executive_summary()
        self._add_major_funding()
        self._add_appointments()
        self._add_education()
        self._add_research()
        self._add_selected_publications(count=10)
        self._add_awards()
        self._add_funding()
        self._add_professional_activities()
        self._add_invited_talks()
        self._add_teaching()
        self._add_affiliations()
        self._add_students()
        self._add_publications()
        self._add_footer()
        print("Saving DOCX...")
        self.doc.save(output_path)
        print(f"✓ DOCX CV generated: {output_path}")


# ----------------------------------------------------------------------------
# Entry point
# ----------------------------------------------------------------------------


def main():
    script_dir = os.path.dirname(os.path.abspath(__file__))
    project_root = os.path.dirname(script_dir)
    yaml_path = os.path.join(project_root, '_data', 'rafael.yml')

    if not os.path.exists(yaml_path):
        print(f"Error: YAML file not found at {yaml_path}")
        return

    output_dir = os.path.join(project_root, 'files', 'cv')
    pdf_path = os.path.join(output_dir, 'RafaelFerreiraDaSilva-cv.pdf')
    docx_path = os.path.join(output_dir, 'RafaelFerreiraDaSilva-cv.docx')

    print(f"Generating CV from {yaml_path}...")

    pdf_gen = CVGenerator(yaml_path, project_root)
    pdf_gen.generate(pdf_path)
    print(f"  PDF size: {os.path.getsize(pdf_path) / 1024:.1f} KB")

    if DOCX_AVAILABLE:
        docx_gen = CVDocxGenerator(yaml_path, project_root)
        docx_gen.generate(docx_path)
        print(f"  DOCX size: {os.path.getsize(docx_path) / 1024:.1f} KB")
    else:
        print("DOCX generation skipped (python-docx not installed)")


if __name__ == '__main__':
    main()
